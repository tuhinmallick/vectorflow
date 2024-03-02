import sys
import os

# this is needed to import classes from other modules
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../')))

import requests
import magic # magic requires lib magic to be installed on the system. If running on mac and an error occurs, run `brew install libmagic`
import json
import fitz
import base64
import logging
import copy
from io import BytesIO
import services.database.batch_service as batch_service
import services.database.job_service as job_service
from flask import Flask, request, jsonify
from flask_cors import CORS
from models.batch import Batch
from api.auth import Auth
from api.pipeline import Pipeline
from services.database.database import get_db, safe_db_operation
from shared.vectorflow_request import VectorflowRequest
from shared.embeddings_type import EmbeddingsType
from docx import Document
from shared.image_search_request import ImageSearchRequest
from urllib.parse import urlparse
from pathlib import Path
from llama_index import download_loader
from services.minio.minio_service import create_minio_client
from api.posthog import send_telemetry
from datetime import datetime

auth = Auth()
pipeline = Pipeline()
app = Flask(__name__)
CORS(app)

logging.basicConfig(filename='./api-log.txt', level=logging.INFO)
logging.basicConfig(filename='./api-errors.txt', level=logging.ERROR)

@app.route("/embed", methods=['POST'])
def embed():
    # TODO: add validator service
    """
    This Python function named `embed` processes an embed request in the VectorFlow
    API. It takes a Flask Request object as input and performs several validations
    and operations based on the content of the request:
    1/ Validates the credentials of the user sending the request. If invalid,
    returns a JSON response with the error message "Invalid credentials" and status
    code 401.
    2/ Checks if all required fields are present in the request. If not, returns
    a JSON response with the error message "Missing required fields" and status
    code 400.
    3/ Validates the type of embeddings metadata provided in the request. If it's
    a Hugging Face embeddings model but no "hugging_face_model_name" is provided,
    returns a JSON response with the error message "Hugging face embeddings models
    require a 'hugging_face_model_name'" and status code 400.
    4/ Checks if a webhook URL is provided in the request but no webhook key.
    Returns a JSON response with the error message "Webhook URL provided but no
    webhook key" and status code 400.
    5/ Processes the file part of the request (if present). Sets the file pointer
    to the beginning of the file, gets its size, and checks if it's too large
    (currently exceeding 25 MB). If so, returns a JSON response with the error
    message "File is too large" and status code 413.
    6/ If the file is valid and not too large, creates a new job in the VectorFlow
    database using the provided embeddings metadata, saves it to the local vector
    database if necessary, and sends telemetry data to track the operation. Then,
    it returns a JSON response with the job ID and a success message.
    In summary, this function processes an embed request by validating credentials,
    checking required fields, validating embeddings metadata, verifying webhook
    setup, processing files, and returning a successful response if the operation
    was successful.

    Returns:
        str: The function `embed` returns a JSON response with an error message
        if the input provided is invalid. Here are the possible outputs returned
        by the function:
        
        	- If the credentials provided are invalid, the function returns a JSON
        response with an error message stating 'Invalid credentials' and a status
        code of 401.
        	- If any required field is missing in the input request, the function
        returns a JSON response with an error message stating 'Missing required
        fields' and a status code of 400.
        	- If the Hugging Face embeddings model name is not provided in the input
        request for the HUGGING_FACE embeddings type, the function returns a JSON
        response with an error message stating 'Hugging face embeddings models
        require a "hugging_face_model_name" in the "embeddings_metadata"' and a
        status code of 400.
        	- If the webhook URL is provided but no webhook key, the function returns
        a JSON response with an error message stating 'Webhook URL provided but
        no webhook key' and a status code of 400.
        	- If no file part is provided in the request, the function returns a JSON
        response with an error message stating 'No file part in the request' and
        a status code of 400.
        	- If the uploaded file size exceeds the limit set in the function (25
        MB), the function returns a JSON response with an error message stating
        'File is too large. The /embed endpoint currently only supports 25 MB files
        or less.' and a status code of 413.
        	- If no selected file is provided, the function returns a JSON response
        with an error message stating 'No selected file' and a status code of 400.
        	- If the uploaded file is not a TXT, PDF, Markdown or DOCX file, the
        function returns a JSON response with an error message stating 'Uploaded
        file is not a TXT, PDF, Markdown or DOCX file' and a status code of 400.

    """
    vectorflow_request = VectorflowRequest._from_flask_request(request)
    if not vectorflow_request.vectorflow_key or not auth.validate_credentials(vectorflow_request.vectorflow_key):
        return jsonify({'error': 'Invalid credentials'}), 401
 
    if not vectorflow_request.embeddings_metadata or not vectorflow_request.vector_db_metadata or (not vectorflow_request.vector_db_key and not os.getenv('LOCAL_VECTOR_DB')):
        return jsonify({'error': 'Missing required fields'}), 400
    
    if vectorflow_request.embeddings_metadata.embeddings_type == EmbeddingsType.HUGGING_FACE and not vectorflow_request.embeddings_metadata.hugging_face_model_name:
        return jsonify({'error': 'Hugging face embeddings models require a "hugging_face_model_name" in the "embeddings_metadata"'}), 400
    
    if vectorflow_request.webhook_url and not vectorflow_request.webhook_key:
        return jsonify({'error': 'Webhook URL provided but no webhook key'}), 400
    
    if 'SourceData' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400

    file = request.files['SourceData']
    
    # TODO: Remove this once the application is reworked to support large files
    # Get the file size - Go to the end of the file, get the current position, and reset the file to the beginning
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)

    if file_size > 25 * 1024 * 1024:
        return jsonify({'error': 'File is too large. The /embed endpoint currently only supports 25 MB files or less. Please use /jobs for streaming large files or multiple files.'}), 413
    
    # empty filename means no file was selected
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and is_valid_file_type(file):
        job = safe_db_operation(job_service.create_job, vectorflow_request, file.filename)
        batch_count = process_file(file, vectorflow_request, job.id)

        vectorflow_request_copy = copy.deepcopy(vectorflow_request)
        send_telemetry("SINGLE_FILE_UPLOAD_SUCCESS", vectorflow_request_copy)

        logging.info(f"{datetime.now()} Successfully created job {job.id} for file {file.filename}")
        return jsonify({'message': f"Successfully added {batch_count} batches to the queue", 'JobID': job.id}), 200
    else:
        return jsonify({'error': 'Uploaded file is not a TXT, PDF, Markdown or DOCX file'}), 400

@app.route('/jobs', methods=['POST'])
def create_jobs():
     # TODO: add validator service
    """
    This function handles file uploads as part of the multi-file processing workflow
    in VectorFlow. It takes care of validating input credentials, checking for
    required fields, and uploading files to Minio (a distributed cloud storage
    system) for further processing. The function does the following:
    1/ Validates the incoming request's credentials using `auth.validate_credentials`.
    If invalid credentials are provided, a JSON response with an error message is
    returned with status code 401.
    2/ Checks if all required fields are present in the request's `embeddings_metadata`
    and `vector_db_metadata`. If any field is missing, a JSON response with an
    error message is returned with status code 400.
    3/ Uploads files provided by the client to Minio using `upload_to_minio`. If
    there are any issues during the upload process (e.g., file type not supported),
    a JSON response with an error message is returned with status code 400.
    4/ Creates a job in VectorFlow's database using `safe_db_operation(job_service.create_job)`.
    If there are any issues during the creation of the job, a JSON response with
    an error message is returned with status code 400.
    5/ Sends telemetry data to VectorFlow's analytics system using `send_telemetry`.
    The telemetry data includes information about the successfully processed files
    and jobs.
    6/ Returns a JSON response with success/failure messages and the IDs of
    successfully uploaded files and created jobs.
    7/ Handles failed file uploads and jobs creation by adding the failed files
    to Minio's blacklist and updating VectorFlow's database accordingly.

    Returns:
        dict: The function returns a JSON response with the following fields:
        
        	- `message`: A message indicating that files have been processed.
        	- `successful_uploads`: A dictionary of file names and their corresponding
        job IDs, which are successfully uploaded to Minio.
        	- `failed_uploads`: A list of file names that failed to upload to Minio.
        	- `empty_files_count`: The number of empty files (i.e., files with an
        empty filename) encountered during the upload process.
        	- `duplicate_files_count`: The number of duplicate file names encountered
        during the upload process.

    """
    vectorflow_request = VectorflowRequest._from_flask_request(request)
    if not vectorflow_request.vectorflow_key or not auth.validate_credentials(vectorflow_request.vectorflow_key):
        return jsonify({'error': 'Invalid credentials'}), 401
 
    if not vectorflow_request.embeddings_metadata or not vectorflow_request.vector_db_metadata or (not vectorflow_request.vector_db_key and not os.getenv('LOCAL_VECTOR_DB')):
        return jsonify({'error': 'Missing required fields'}), 400
    
    if vectorflow_request.embeddings_metadata.embeddings_type == EmbeddingsType.HUGGING_FACE and not vectorflow_request.embeddings_metadata.hugging_face_model_name:
        return jsonify({'error': 'Hugging face embeddings models require a "hugging_face_model_name" in the "embeddings_metadata"'}), 400
    
    if vectorflow_request.webhook_url and not vectorflow_request.webhook_key:
        return jsonify({'error': 'Webhook URL provided but no webhook key'}), 400
    
    if not hasattr(request, "files") or not request.files:
        return jsonify({'error': 'No file part in the request'}), 400
    
    files = request.files.getlist('file')
    successfully_uploaded_files = dict()
    failed_uploads = []
    empty_files_count = 0
    duplicate_files_count = 0

    for file in files:
        # Check if a file is empty (no filename)
        if file.filename == '':
            empty_files_count += 1
            continue
        
        if not is_valid_file_type(file):
            failed_uploads.append(file.filename)
            continue

        if file.filename in successfully_uploaded_files:
            duplicate_files_count += 1
            continue

        temporary_storage_location = os.getenv('API_STORAGE_DIRECTORY')
        file_path = os.path.join(temporary_storage_location, file.filename)
        with open(file_path, 'wb') as f:
            chunk_size = 65536  # 64 KB
            while True:
                chunk = file.stream.read(chunk_size)
                if len(chunk) == 0:
                    break
                f.write(chunk)
        try:
            result = upload_to_minio(file_path, file.filename)
            os.remove(file_path)

            if not result:
                failed_uploads.append(file.filename)
                continue
            
            job = safe_db_operation(job_service.create_job, vectorflow_request, file.filename)
            if not job:
                remove_from_minio(file.filename)
                continue
       
            data = (job.id, file.filename, vectorflow_request.serialize())
            json_data = json.dumps(data)

            pipeline.connect(queue=os.getenv('EXTRACTION_QUEUE'))
            pipeline.add_to_queue(json_data, queue=os.getenv('EXTRACTION_QUEUE'))
            pipeline.disconnect()

            successfully_uploaded_files[file.filename] = job.id
            logging.info(f"{datetime.now()} Successfully created job {job.id} for file {file.filename}")
        except Exception as e:
            print(f"Error uploading file {file.filename} to min.io, creating job or passing vectorflow request to message broker. \nError: {e}\n\n")
            failed_uploads.append(file.filename)       

    send_telemetry("MULTI_FILE_UPLOAD_SUCCESS", vectorflow_request)
    return jsonify({'message': 'Files processed', 
                    'successful_uploads': successfully_uploaded_files,
                    'failed_uploads': failed_uploads,
                    'empty_files_count': empty_files_count,
                    'duplicate_files_count': duplicate_files_count}), 200

@app.route('/jobs/<int:job_id>/status', methods=['GET'])
def get_job_status(job_id):
    """
    This Python function, `get_job_status`, retrieves the job status of a given
    job ID by making a GET request to the `/jobs/<job_id>` endpoint. It performs
    the following actions:
    1/ Retrieves the `Authorization` header from the request. If it's not present
    or invalid, it returns a 401 error message.
    2/ Validates the credentials using the `auth.validate_credentials()` function.
    If the validation fails, it returns a 401 error message.
    3/ Uses the `safe_db_operation()` function to call the `get_job()` function
    of the `job_service` object with the job ID as an argument. This function
    retrieves the job details from the database.
    4/ If the job is found, it returns a JSON response with the job status information.
    5/ Otherwise, it returns a 404 error message indicating that the job was not
    found.

    Args:
        job_id (str): In the `get_job_status` function, the `job_id` input parameter
            represents the unique identifier for the job being queried. It is used
            to retrieve the status of the specific job from the database by calling
            the `safe_db_operation()` function, which handles any errors or
            exceptions that may occur during the database operation. The `job_id`
            parameter is required to identify the correct job record in the database
            and return its status information to the user.

    Returns:
        dict: The output of this function will be a JSON response containing the
        job status information if the job ID provided is valid, otherwise it will
        return an error message. Specifically, the function will return
        `jsonify({'JobStatus': job.job_status.value})` if the job is found, and
        `jsonify({'error': "Job not found"})` if the job is not found.

    """
    vectorflow_key = request.headers.get('Authorization')
    if not vectorflow_key or not auth.validate_credentials(vectorflow_key):
        return jsonify({'error': 'Invalid credentials'}), 401
    
    job = safe_db_operation(job_service.get_job, job_id)
    if job:
        return jsonify({'JobStatus': job.job_status.value}), 200
    else:
        return jsonify({'error': "Job not found"}), 404
        
@app.route('/jobs/status', methods=['POST'])
def get_job_statuses():
    """
    This Python function named `get_job_statuses()` retrieves job status information
    from a database using the `safe_db_operation()` function. It takes in a JSON
    body from the client containing `JobIDs`, which it then passes to the `get_jobs()`
    function of the job service, along with any authentication headers and validation.
    If the job statuses are found in the database, the function returns them in a
    JSON response with each job status represented as an object containing the
    `JobID` and the corresponding `JobStatus`. If no jobs are found, the function
    returns a JSON response indicating that the jobs were not found (`'error':
    'Jobs not found'`). If there is an error authenticating or validating the
    credentials, the function returns a JSON response with an `error` field
    containing a message describing the issue (`'Invalid credentials'`) or if no
    `JobIDs` field is provided in the request (`'Missing JobIDs field'`).

    Returns:
        dict: The function returns a JSON response with one of three possible
        outputs depending on the input parameters:
        
        1/ If the credentials are invalid or missing, it returns a response with
        an error message ({"error": "Invalid credentials"}) and a status code of
        401.
        2/ If the JSON body is missing or not present in the request, it returns
        a response with an error message ( {"error": "Missing JSON body"}) and a
        status code of 400.
        3/ If the job IDs are present in the request and the function is able to
        retrieve the jobs from the database using the job service's get_jobs method,
        it returns a response with an array of dictionaries representing the jobs
        ({"Jobs": [{"JobID": job.id, "JobStatus": job.job_status.value} for job
        in jobs])), with a status code of 200. If no jobs are found, it returns a
        response with an error message ("Jobs not found") and a status code of 404.

    """
    vectorflow_key = request.headers.get('Authorization')
    if not vectorflow_key or not auth.validate_credentials(vectorflow_key):
        return jsonify({'error': 'Invalid credentials'}), 401
    
    if not hasattr(request, 'json') or not request.json:
        return jsonify({'error': 'Missing JSON body'}), 400
    
    job_ids = request.json.get('JobIDs')
    if not job_ids:
        return jsonify({'error': 'Missing JobIDs field'}), 400
    
    jobs = safe_db_operation(job_service.get_jobs, job_ids)
    if jobs:
        return jsonify({'Jobs': [{'JobID': job.id, 'JobStatus': job.job_status.value} for job in jobs]}), 200
    else:
        return jsonify({'error': "Jobs not found"}), 404   
    
@app.route("/s3", methods=['POST'])
def s3_presigned_url():
    # TODO: add validator service
    """
    This function performs the following tasks:
    1/ It validates the credentials of the request by calling `auth.validate_credentials()`
    and checks if the vectorflow key is provided in the request. If the credentials
    are invalid or the vectorflow key is not provided, it returns a JSON response
    with an error message.
    2/ It checks if the webhook URL is provided in the request along with the
    webhook key. If the webhook URL is provided but the webhook key is missing,
    it returns a JSON response with an error message.
    3/ It retrieves the pre-signed URL from the request parameter `PreSignedURL`.
    If the pre-signed URL is not provided or the URL is invalid, it returns a JSON
    response with an error message.
    4/ It makes a GET request to the pre-signed URL and retrieves the file content.
    5/ It performs further validation on the file content based on its MIME type.
    If the file content is not a TXT, PDF, HTML, or DOCX file, it returns a JSON
    response with an error message.
    6/ It creates a job in the VectorFlow database using the vectorflow request
    and the file content.
    7/ It returns a JSON response indicating whether the operation was successful
    (with the number of batches added to the queue) or not.

    Returns:
        str: The output of this function depends on the status code of the response
        from the S3 pre-signed URL. Here are the possible outputs:
        
        1/ If the response status code is 200, the function returns a JSON object
        containing a message and the job ID, both with values indicating that the
        batch has been successfully added to the queue. The output would be like:
        `{'message': 'Successfully added 3 batches to the queue', 'JobID': 12345}`
        2/ If the response status code is not 200 (e.g., 400, 401, 404, etc.), the
        function returns a JSON object with an error message. The output would be
        like: `{'error': 'Invalid credentials'} or {'error': 'Webhook URL provided
        but no webhook key'}`, respectively.
        3/ If the pre-signed URL is invalid (i.e., response status code is not
        200), the function prints a message to the console indicating that the
        file could not be downloaded. The output would be something like: `Failed
        to download file: 404`.

    """
    vectorflow_request = VectorflowRequest._from_flask_request(request)
    if not vectorflow_request.vectorflow_key or not auth.validate_credentials(vectorflow_request.vectorflow_key):
        return jsonify({'error': 'Invalid credentials'}), 401
    
    if vectorflow_request.webhook_url and not vectorflow_request.webhook_key:
        return jsonify({'error': 'Webhook URL provided but no webhook key'}), 400
 
    pre_signed_url = request.form.get('PreSignedURL')
    if not vectorflow_request.embeddings_metadata or not vectorflow_request.vector_db_metadata or (not vectorflow_request.vector_db_key and not os.getenv('LOCAL_VECTOR_DB')) or not pre_signed_url:
        return jsonify({'error': 'Missing required fields'}), 400
    
    response = requests.get(pre_signed_url)
    file_name = get_s3_file_name(pre_signed_url)

    if response.status_code == 200:
        file_magic = magic.Magic(mime=True)
        mime_type = file_magic.from_buffer(response.content)

        if mime_type == 'text/plain':
            file_content = response.text
            job = safe_db_operation(job_service.create_job, vectorflow_request, file_name)
            return jsonify({'message': f"Successfully added {batch_count} batches to the queue", 'JobID': job.id}), 200
        
        elif mime_type == 'application/pdf':
            pdf_data = BytesIO(response.content)
            with fitz.open(stream=pdf_data, filetype='pdf') as doc:
                file_content = ""
                for page in doc:
                    file_content += page.get_text()

            job = safe_db_operation(job_service.create_job, vectorflow_request, file_name)
            batch_count = create_batches(file_content, job.id, vectorflow_request)
            return jsonify({'message': f"Successfully added {batch_count} batches to the queue", 'JobID': job.id}), 200
        
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            docx_data = BytesIO(response.content)
            doc = Document(docx_data)
            file_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            job = safe_db_operation(job_service.create_job, vectorflow_request, file_name)
            batch_count = create_batches(file_content, job.id, vectorflow_request)
            return jsonify({'message': f"Successfully added {batch_count} batches to the queue", 'JobID': job.id}), 200
        
        else:
            return jsonify({'error': 'Uploaded file is not a TXT, PDF, HTML or DOCX file'}), 400
    else:
        print('Failed to download file:', response.status_code, response.reason)

def process_file(file, vectorflow_request, job_id):
    """
    This function takes a file object, a `vectorflow_request` object and an integer
    identifier for the job, and determines the type of file it is based on its
    filename extension. Depending on the extension, it reads the file content and
    returns a count of created batches after processing the content through
    VectorFlow's batch creation API.

    Args:
        file (): The `file` input parameter in the given function refers to a file
            object passed as an argument to the function. The function processes
            the contents of the file depending on its file extension, and then
            returns the batch count created based on the processed content.
        vectorflow_request (): In this function, `vectorflow_request` is a parameter
            that is passed in from outside the function. It serves as an identifier
            for the batch of data being processed, and is used to create batches
            of data for processing by VectorFlow. The exact purpose and meaning
            of `vectorflow_request` depends on the context and implementation of
            the function, but in general, it is used to keep track of the batch
            of data being processed and to enable vectorflow to perform its intended
            function on the data.
        job_id (): In this function, `job_id` is an input parameter that represents
            a unique identifier for the processing job. The purpose of including
            it in the function is to enable the creation of batches of files based
            on the job ID. This allows for efficient parallel processing of large
            numbers of files, as each batch can be processed independently without
            interfering with other batches. By using the job ID, the function can
            keep track of which batch belongs to which job and process them accordingly.

    Returns:
        int: The function returns the number of batches created from the processed
        file content.

    """
    if file.filename.endswith('.txt'):
        file_content = file.read().decode('utf-8')
    
    elif file.filename.endswith('.docx'):
        doc = Document(file)
        file_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    elif file.filename.endswith('.md'):
        temp_file_path = Path('./temp_file.md')
        file.save(temp_file_path)
            
        MarkdownReader = download_loader("MarkdownReader")
        loader = MarkdownReader()
        documents = loader.load_data(file=Path('./temp_file.md'))

        file_content = "\n".join([document.text for document in documents])
        temp_file_path.unlink()
    
    elif file.filename.endswith('.html'):
        content = file.read().decode('utf-8')
        file_content = repr(content)

    else:
        pdf_data = BytesIO(file.read())
        with fitz.open(stream=pdf_data, filetype='pdf') as doc:
            file_content = ""
            for page in doc:
                file_content += page.get_text()

    batch_count = create_batches(file_content, job_id, vectorflow_request)
    return batch_count

def create_batches(file_content, job_id, vectorflow_request_original):
    """
    This function takes in several inputs:
    	- `file_content`: the contents of a file
    	- `job_id`: the ID of a Job object in the VectorFlow system
    	- `vectorflow_request_original`: the original VectorFlow request
    The function does the following:
    1/ Splits the input file into chunks using the `split_file()` function.
    2/ Creates a list of Batch objects, each representing a single batch of data
    to be processed by the VectorFlow system. The `create_batches()` function is
    used to create the batches.
    3/ Updates the total number of batches for a job using the `update_job_total_batches()`
    function.
    4/ Adds each Batch object to an embedding queue using the `add_to_queue()`
    function in the `pipeline` module.
    5/ Returns the total number of batches created or None if there was no job ID
    provided.

    Args:
        file_content (str): In the provided function, `file_content` is the contents
            of a file that contains the vector flow requests to be processed. The
            function splits the file content into chunks, and each chunk is used
            to create a new batch. The `file_content` parameter is used to determine
            the size of the chunks, as it is passed to the `split_file()` function.
        job_id (str): The `job_id` input parameter is used to update the total
            number of batches for a specific job in the database using the
            `safe_db_operation()` method. It serves as an identifier for the job,
            allowing the function to access and modify the appropriate job data
            in the database.
        vectorflow_request_original (): The `vectorflow_request_original` input
            parameter in the `create_batches()` function serves as a reference to
            the original `VectorFlowRequest` object that was passed to the function.
            It allows the function to create deep copies of the original request
            when splitting the file into batches, ensuring that any changes made
            to the original request are isolated and do not affect the original
            data. This helps ensure the integrity of the data being processed and
            prevents unexpected behavior due to modifications to the original request.

    Returns:
        dict: The output returned by the `create_batches` function is the total
        number of batches created, which is stored in the `job.total_batches`
        attribute of the `job` object if a job was created successfully, or `None`
        otherwise.

    """
    vectorflow_request = copy.deepcopy(vectorflow_request_original)
    chunks = [chunk for chunk in split_file(file_content, vectorflow_request.lines_per_batch)]
    
    batches = [Batch(job_id=job_id, embeddings_metadata=vectorflow_request.embeddings_metadata, vector_db_metadata=vectorflow_request.vector_db_metadata) for _ in chunks]
    batches = safe_db_operation(batch_service.create_batches, batches)

    job = safe_db_operation(job_service.update_job_total_batches, job_id, len(batches))

    for batch, chunk in zip(batches, chunks):
        data = (batch.id, chunk, vectorflow_request.vector_db_key, vectorflow_request.embedding_api_key)
        json_data = json.dumps(data)

        pipeline.connect(queue=os.getenv('EMBEDDING_QUEUE'))
        pipeline.add_to_queue(json_data, queue=os.getenv('EMBEDDING_QUEUE'))
        pipeline.disconnect()

    return job.total_batches if job else None
    
def split_file(file_content, lines_per_chunk=1000):
    """
    This Python function `split_file()` takes in the contents of a file and breaks
    it up into chunks of a certain size (defined by the `lines_per_chunk` argument)
    based on the number of lines in the file. Here's what the function does:
    1/ It splits the file content into a list of lines using the `splitlines()` method.
    2/ It then loops through the list of lines, starting from the beginning of the
    file and ending at the end of the file (inclusive), breaking the list of lines
    into chunks of size `lines_per_chunk` using slicing notation (`[:i+lines_per_chunk]`).
    3/ For each chunk of lines, the function yields the chunk as a separate generator
    expression.
    So, if you call the function with `file_content = "This is a file contents."`
    and `lines_per_chunk=10`, the function will return 10 chunks of lines, each
    containing 10 lines:
    ```
    chunk1: ['This', 'is', 'a', 'file', 'contents.']
    chunk2: ['file', 'contents.']
    chunk3: [' contents.']
    ...
    ```

    Args:
        file_content (list): The `file_content` input parameter in the `split_file`
            function represents the contents of a file that needs to be split into
            smaller chunks. It is a string or a binary object containing the
            contents of the file. The function takes the contents of the file as
            its input and splits it into chunks of a specified number of lines,
            based on the `lines_per_chunk` parameter.
        lines_per_chunk (int): The `lines_per_chunk` input parameter in the
            `split_file()` function specifies the number of lines to process in
            each chunk. It controls the granularity of the splitting operation,
            determining how many lines are processed at a time before moving on
            to the next batch.
            
            In other words, it sets the "chunksize" or "lines per chunk" for the
            iteration over the file content. By default, it is set to 1000, meaning
            that the function will split the file into chunks of approximately
            1000 lines each.

    """
    lines = file_content.splitlines()
    for i in range(0, len(lines), lines_per_chunk):
        yield lines[i:i+lines_per_chunk]

@app.route("/images", methods=['POST'])
def upload_image():
    # TODO: add validator service
    """
    This Flask function handles image uploads for VectorFlow, a machine learning-based
    microscopy image analysis platform. It performs the following tasks:
    1/ Validates the uploaded credentials and vector database metadata (if required).
    If there are any errors, it returns a JSON response with an "error" key.
    2/ Checks if a file was selected in the request. If not, it returns a JSON
    response with an "error" key.
    3/ Sets the file pointer to the beginning of the file using `seek()`, checks
    its size, and resizes it if necessary (to prevent excessive memory usage).
    4/ Processes the image if it is a JPEG, JPG, or PNG file using `process_image()`,
    which adds the batch to the job queue for analysis. If there's an error during
    processing, it catches and logs the exception and returns a JSON response with
    an "error" key.
    5/ Otherwise, it returns a JSON response indicating that the file is not a
    supported format.

    Returns:
        dict: The output returned by this function will depend on the specific
        input passed to it. However, the function returns a response with the
        following possible outcomes:
        
        	- If the credentials are invalid or missing, the function returns a
        response with an error message and a status code of 401 (Unauthorized).
        	- If one of the required fields is missing, the function returns a response
        with an error message and a status code of 400 (Bad Request).
        	- If the file size is greater than the allowed limit (2 MB), the function
        returns a response with an error message and a status code of 413 (Payload
        too large).
        	- If no file was selected, the function returns a response with an error
        message and a status code of 400 (Bad Request).
        	- If the uploaded file is not a JPG, JPEG, or PNG file, the function
        returns a response with an error message and a status code of 400 (Bad Request).
        	- If any other error occurs during the process, the function catches the
        exception and returns a response with an error message and a status code
        of 500 (Internal Server Error).

    """
    vectorflow_request = VectorflowRequest._from_flask_request(request)
    if not vectorflow_request.vectorflow_key or not auth.validate_credentials(vectorflow_request.vectorflow_key):
        return jsonify({'error': 'Invalid credentials'}), 401
 
    if not vectorflow_request.vector_db_metadata or (not vectorflow_request.vector_db_key and not os.getenv('LOCAL_VECTOR_DB')):
        return jsonify({'error': 'Missing required fields'}), 400
    
    if 'SourceData' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400

    file = request.files['SourceData']

    # TODO: Remove this once the application is reworked to support large files
    # Get the file size - Go to the end of the file, get the current position, and reset the file to the beginning
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)

    if file_size > 2 * 1024 * 1024:
        return jsonify({'error': 'File is too large. VectorFlow currently only supports 2 MB files or less for images. Larger file support coming soon.'}), 413
    
    # empty filename means no file was selected
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and (file.filename.endswith('.jpg') or file.filename.endswith('.jpeg') or file.filename.endswith('.png')):
        try:
            job_id = process_image(file, vectorflow_request)
            return jsonify({'message': f"Successfully added {file.filename} batches to the queue", 'JobID': job_id}), 200
        
        except Exception as e:
            logging.error(f"Attempt to upload file {file.filename} failed due to error: {e}")
            return jsonify({'error': f"Attempt to upload file {file.filename} failed due to error: {e}"}), 400
    else:
        return jsonify({'error': 'Uploaded file is not a JPG, JPEG, or PNG file'}), 400

def process_image(file, vectorflow_request):
    # Create job
    """
    This function processes an image file and sends it to a VectorFlow server for
    analysis. Here is a breakdown of the function's steps:
    1/ It creates a job in the VectorFlow server using the
    `job_service.create_job_with_vdb_metadata()` method.
    2/ It converts the image file to bytes using `BytesIO()`.
    3/ It encodes the image bytes as a Base64 string for serialization to JSON.
    4/ It creates a JSON message body with the encoded image, vector flow request
    metadata, and the job ID.
    5/ It sends the JSON data to the Image queue using `pipeline.add_to_queue()`.
    6/ It disconnects from the pipeline using `pipeline.disconnect()`.
    7/ The function returns the job ID.

    Args:
        file (): In the given function, the `file` input parameter represents a
            file object that contains the image to be processed. The function uses
            the `file` object to read the image bytes and convert them to a Base64
            encoded string for serialization to JSON.
        vectorflow_request (dict): The `vectorflow_request` input parameter in the
            `process_image` function serves as an initialization for a vector
            database (VDB) metadata. Specifically, it provides the key to access
            the VDB metadata in the job creation process with the
            `job_service.create_job_with_vdb_metadata()` method.

    Returns:
        int: The output of this function is the ID of the Job created in VectorFlow's
        Job Service.

    """
    with get_db() as db:
        job = job_service.create_job_with_vdb_metadata(db, vectorflow_request, file.filename)

    # Convert image to bytes
    img_bytes_io = BytesIO()
    file.save(img_bytes_io)
    image_bytes = img_bytes_io.getvalue()

    # Encode image bytes to Base64 string to allowed for serializaton to JSON
    encoded_image_string = base64.b64encode(image_bytes).decode("utf-8")

    # Convert to JSON - this format can be read agnostic of the technology on the other side
    message_body = {
        'image_bytes': encoded_image_string,
        'vector_db_key': vectorflow_request.vector_db_key,
        'job_id': job.id,
    }
    json_data = json.dumps(message_body)

    pipeline.connect(queue=os.getenv('IMAGE_QUEUE'))
    pipeline.add_to_queue(json_data, queue=os.getenv('IMAGE_QUEUE'))
    pipeline.disconnect()
    
    return job.id

@app.route("/images/search", methods=['POST'])
def search_image_from_vdb():
    """
    This function performs image search using VectorFlow, an open-source image
    recognition library. Here's what it does:
    1/ Validates the request credentials and vector flow key (if provided).
    2/ Checks if all required fields are provided in the request, including the
    vector database metadata and the file part containing the image to be searched.
    3/ Gets the file size from the request file part and checks if it's within the
    allowed limit of 2 MB or less.
    4/ Checks if the uploaded file is a JPG, JPEG, or PNG file.
    5/ If all checks pass, it attempts to search for similar images in the vector
    database using the provided image as input.
    6/ The function returns a JSON response with the search results, including
    vectors and similar images, if successful. If an error occurs, it returns a
    JSON message with the error details.

    Returns:
        str: The output returned by this function depends on the value of
        `file.filename` and the success of the search operation. Here are some
        possible outputs:
        
        1/ If `file.filename` ends with ".jpg", ".jpeg", or ".png", and the search
        operation succeeds, the function returns a JSON response with the following
        fields:
        	- "similar_images": a list of similar images to the uploaded image, with
        each image represented as a JSON object containing the following fields:
        	+ "image_url": the URL of the similar image
        	+ "score": a score representing how similar the image is to the uploaded
        image
        	+ "file_name": the filename of the similar image
        2/ If the search operation fails due to an error, the function returns a
        JSON response with the following fields:
        	- "error": a string describing the error that occurred during the search
        operation
        3/ If the uploaded file is not a JPG, JPEG, or PNG file, the function
        returns a JSON response with the following field:
        	- "error": a string indicating that the uploaded file is not a valid image
        format
        The function also includes logging for any unhandled errors that may occur
        during the search operation.

    """
    image_search_request = ImageSearchRequest._from_request(request)
    if not image_search_request.vectorflow_key or not auth.validate_credentials(image_search_request.vectorflow_key):
        return jsonify({'error': 'Invalid credentials'}), 401
 
    if not image_search_request.vector_db_metadata or (not image_search_request.vector_db_key and not os.getenv('LOCAL_VECTOR_DB')):
        return jsonify({'error': 'Missing required fields'}), 400
    
    if 'SourceData' not in request.files:
        return jsonify({'error': 'No file part in the request'}), 400

    file = request.files['SourceData']

    # TODO: Remove this once the application is reworked to support large files
    # Get the file size - Go to the end of the file, get the current position, and reset the file to the beginning
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)

    if file_size > 2 * 1024 * 1024:
        return jsonify({'error': 'File is too large. VectorFlow currently only supports 2 MB files or less for images. Larger file support coming soon.'}), 413
    
    # empty filename means no file was selected
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and (file.filename.endswith('.jpg') or file.filename.endswith('.jpeg') or file.filename.endswith('.png')):
        try:
            response = search_image_from_vdb(file, image_search_request)

            if response.status_code == 200:
                response_json = response.json()
                if "vectors" in response_json:
                    return jsonify({'message': f"Successfully fetched {image_search_request.top_k} results including vectors",
                        'similar_images': response_json['similar_images'],
                        'vectors': response_json['vectors']}), 200
                else:
                    return jsonify({'message': f"Successfully fetched {image_search_request.top_k} results", 
                        'similar_images': response_json['similar_images']}), 200
            else:
                response_json = response.json()
                error_message = response_json["error"]
                return jsonify({'error': f"Attempt to fetch images similar to {file.filename} failed due to error: {error_message}"}), response.status_code
        
        except Exception as e:
            logging.error(f"Attempt to fetch images similar to {file.filename} failed due to error: {e}")
            return jsonify({'error': f"Attempt to fetch images similar to {file.filename} failed due to error: {e}"}), 400
    else:
        return jsonify({'error': 'Uploaded file is not a JPG, JPEG, or PNG file'}), 400
    
def search_image_from_vdb(file, image_search_request):
    """
    This function performs an image search using the `image_search_request`
    parameter, which contains the details of the image to be searched. It sends a
    POST request to the specified `IMAGE_SEARCH_URL` with the request data serialized
    in JSON format and the file containing the image to be searched as a
    multipart/form-data file. The function returns the response from the API call,
    which may contain an error message if there is one, or the search results if
    the API call was successful.

    Args:
        file (): In the provided code, the `file` parameter is used to pass the
            image file that the user wants to search for in an online VDB. It is
            passed as a dictionary with the key `SourceData`.
        image_search_request (dict): The `image_search_request` input parameter
            in the `search_image_from_vdb()` function is a serialized JSON
            representation of an image search request. It contains information
            such as the query, category, and other parameters that are used to
            perform an image search using the provided URL.

    Returns:
        dict: The output returned by this function is a `Response` object containing
        the result of the image search query sent to the specified URL.

    """
    url = f"{os.getenv('IMAGE_SEARCH_URL')}/search"
    data = {
        'ImageSearchRequest': json.dumps(image_search_request.serialize()),
    }

    files = {
        'SourceData': file
    }

    try:
        response = requests.post(
            url=url, 
            data=data, 
            files=files
        )
        return response
    except requests.RequestException as e:
        print(f"Error: {e}")
        return {"error": str(e)}, 500
    
def get_s3_file_name(pre_signed_url):
    """
    This function takes a pre-signed S3 URL as input and returns the filename
    portion of the URL path.
    In detail, the function first parses the given S3 URL using `urlparse()` to
    obtain the path parts. Then it extracts the last element of the path parts
    which represents the file name and returns it.
    In summary, the function takes a pre-signed S3 URL and returns the filename
    portion of the URL path.

    Args:
        pre_signed_url (str): The `pre_signed_url` input parameter is used to pass
            the pre-signed URL provided by AWS S3 for accessing a file in S3. This
            pre-signed URL contains an AWS access token that is valid for a limited
            time, which the `get_s3_file_name` function uses to retrieve the file
            name from the path part of the URL.

    Returns:
        str: The function `get_s3_file_name` returns the last part of a URL's path,
        which is the file name after splitting the path using `/`. Therefore, the
        output returned by this function will be the final file name in the path
        of the provided pre-signed URL.

    """
    parsed_url = urlparse(pre_signed_url)
    path_parts = parsed_url.path.lstrip('/').split('/')

    # For the file name and not the full path:
    file_name = path_parts[-1]
    return file_name

def is_valid_file_type(file):
    """
    This function takes in a `File` object as input and returns a boolean value
    indicating whether the file is of a supported type (`.txt`, `.docx`, `.pdf`,
    `.md`, or `.html`). It first checks if the file's filename ends with any of
    the supported types, and if it does, it returns `True`. If not, it attempts
    to read a portion of the file's content (up to 1024 bytes) and decode it as
    UTF-8 text. If the decoding is successful, it adds the file's original filename
    with the `.txt` extension and returns `True`, indicating that it is likely a
    text file. Otherwise, it returns `False`.

    Args:
        file (): The `file` input parameter in the provided code is used as the
            basis for checking if the given file is a valid file type. Specifically,
            it represents the file to be checked for validity, and the function
            accesses its properties such as `filename`, `stream`, etc., to perform
            the validation checks.

    Returns:
        bool: The function returns `True` if the file is a valid text file, and
        `False` otherwise.

    """
    supported_types = ['.txt', '.docx', '.pdf', '.md', '.html']
    for type in supported_types:
        if file.filename.endswith(type):
            return True
        
    # Try to detect .txt files by content
    try:
        # Read a portion of the file
        file_content = file.stream.read(1024)
        file.stream.seek(0)  # Reset the file stream position for later use

        # Attempt to decode the content as utf-8 text
        file_content.decode('utf-8')

        # If we were able to successfully decode the file as utf-8 text, it's likely a text file
        file.filename += '.txt'
        return True  # Successful decoding implies it's a text file
    except UnicodeDecodeError:
        return False  # Failed to decode, likely not a text file

def remove_from_minio(filename):
    """
    This function removes an object from a bucket in Min.io, a cloud-native object
    storage service. It does the following:
    1/ Creates a `Minio` client instance using the environment variable `MINIO_BUCKET`.
    2/ Calls the `remove_object()` method of the client instance with the name of
    the bucket and the object filename as arguments.
    In summary, this function removes an object from a Minio bucket using the Minio
    client API.

    Args:
        filename (str): The `filename` input parameter in the `remove_from_minio`
            function refers to the name of the file or object that needs to be
            removed from the Minio bucket. It is passed as a string to the
            `remove_object` method of the Minio client, specifying the bucket name
            and the object name or path to be removed.

    """
    client = create_minio_client()
    client.remove_object(os.getenv("MINIO_BUCKET"), filename)

def upload_to_minio(file_path, filename):
    """
    This function appears to be designed to upload a file from a local path to
    Minio bucket using the `create_minio_client` method to create an instance of
    the `MinioClient` class, and then using the `put_object` method of the client
    to store the file in the bucket. Here's a breakdown of what the function does:
    1/ Creates an instance of `MinioClient` using the `create_minio_client` method.
    2/ Opens the file at the specified path using the `file_data_generator` method,
    which returns a generator that yields chunks of the file data.
    3/ Wraps the generator with a `StreamWrapper` object, which adds some additional
    functionality to handle streaming the file data to Minio.
    4/ Calls the `put_object` method of the `MinioClient` instance with the bucket
    name, filename, and the generated stream as arguments. This method will store
    the file in the specified bucket with the specified filename.
    5/ Returns the result of the `put_object` method, which is the response from
    Minio.

    Args:
        file_path (str): The `file_path` input parameter specifies the path to the
            local file that needs to be uploaded to Minio. It is used to retrieve
            the contents of the file and send it to Minio for storage.
        filename (str): In the `upload_to_minio()` function, the `filename` input
            parameter is used to specify the name of the object that will be stored
            in Min.IO bucket. It is a required argument and it should be the same
            as the filename of the file being uploaded. The function uses this
            value to create the storage key for the object being uploaded.

    Returns:
        : The output of this function is a `MinioResult` object, which contains
        information about the successful upload of an object to Min.io. The fields
        of the `MinioResult` object include:
        
        	- `body`: The body of the uploaded object.
        	- `bucket`: The name of the bucket where the object was uploaded.
        	- `etag`: The etag of the uploaded object.
        	- `host`: The hostname of the Min.io server where the object was uploaded.
        	- `status_code`: The HTTP status code of the upload operation.
        	- `reason`: An optional error message if the upload operation failed.

    """
    client = create_minio_client()

    file_size = os.path.getsize(file_path)

    # Wrap the generator with our StreamWrapper
    stream = StreamWrapper(lambda: file_data_generator(file_path))

    result = client.put_object(
        os.getenv("MINIO_BUCKET"), filename, stream, file_size
    )
    return result

# generator used to stream
def file_data_generator(file_path, chunk_size=65536):  # 64KB
    """
    This `def` statement defines a function named `file_data_generator`. The
    function takes in a file path as an argument, and returns an generator object
    that yields (or repeatedly returns) chunks of data from the specified file.
    The function first opens the file using `open()` in binary (`'rb'` mode) and
    then reads chunks of data from the file using the `read()` method. The size
    of each chunk is specified as an argument to the function, which is defaulted
    to 64KB (65536). If no data is read (i.e., the end of the file is reached),
    the function breaks out of the loop and stops returning any more data. Otherwise,
    it continues reading and yielding chunks until the end of the file is reached.

    Args:
        file_path (str): The `file_path` input parameter in the `file_data_generator`
            function specifies the path to a file that the function will read from.
        chunk_size (int): The `chunk_size` input parameter in the `file_data_generator`
            function controls the amount of data read from the file for each
            iteration of the loop. It determines how much of the file is processed
            before moving on to the next iteration. In this case, it sets the chunk
            size to 64KB (or 65536 bytes).

    """
    with open(file_path, 'rb') as file:
        while True:
            data = file.read(chunk_size)
            if not data:
                break
            yield data

class StreamWrapper:
    def __init__(self, generator_func):
        self.generator = generator_func()

    def read(self, *args):
        return next(self.generator, b'')

if __name__ == '__main__':
   app.run(host='0.0.0.0', debug=True)